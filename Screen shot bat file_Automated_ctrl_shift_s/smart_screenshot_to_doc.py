import os
import time
import keyboard
import pyautogui
from docx import Document
from docx.shared import Inches

# 📌 Ask user for Word document name
file_name = input("📝 Enter file name for Word document (without .docx): ").strip()
if not file_name:
    file_name = "Screenshots_Report"

output_file = file_name + ".docx"

# Setup folders
folder = "screenshots"
os.makedirs(folder, exist_ok=True)

doc = Document()
doc.add_heading("Screenshot Report", level=1)
count = 1

print("\n📸 Screenshot tool running...")
print("🧷 Press Ctrl + Shift + S to take a screenshot.")
print("🛑 Press Ctrl + Shift + Q to save and quit.\n")

while True:
    if keyboard.is_pressed('ctrl+shift+s'):
        filename = os.path.join(folder, f"screenshot_{count}.png")
        pyautogui.screenshot(filename)
        doc.add_paragraph(f"Screenshot {count}")
        doc.add_picture(filename, width=Inches(6))

        print(f"✅ Screenshot {count} captured.")
        print(f"📄 Saved to: {filename}\n")

        count += 1
        while keyboard.is_pressed('ctrl+shift+s'):
            time.sleep(0.2)

    if keyboard.is_pressed('ctrl+shift+q'):
        doc.save(output_file)
        print(f"\n📁 Word document saved as: {output_file}")
        print("👋 Exiting...")
        break

    time.sleep(0.1)
