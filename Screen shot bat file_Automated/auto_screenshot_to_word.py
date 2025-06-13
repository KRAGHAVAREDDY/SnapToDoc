import pyautogui
import time
from docx import Document
from docx.shared import Inches
import os
import keyboard  # NEW: to detect keypress

# Prepare document
doc = Document()
doc.add_heading('Screenshots Report', 0)

# Create folder if not exists
folder = "screenshots"
if not os.path.exists(folder):
    os.makedirs(folder)

i = 1
print("Started capturing screenshots... Press 'q' to stop.")

while True:
    if keyboard.is_pressed('q'):
        print("Stopping...")
        break

    filename = f"{folder}/screenshot_{i}.png"
    image = pyautogui.screenshot()
    image.save(filename)

    doc.add_paragraph(f"Screenshot {i}:")
    doc.add_picture(filename, width=Inches(6))
    doc.add_paragraph("")

    print(f"Captured {filename}")
    time.sleep(3)  # Wait for 3 seconds
    i += 1

# Save final document
doc.save("Screenshots_Report.docx")
print("Saved Screenshots_Report.docx")
