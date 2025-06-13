import pyautogui
from docx import Document
from docx.shared import Inches
from pynput import keyboard
import time
import os

# Create or open a Word doc
doc = Document()
doc.add_heading('Manual Delayed Screenshot Report', 0)

folder = "screenshots"
if not os.path.exists(folder):
    os.makedirs(folder)

count = 1
print("Press 's' to take a screenshot after 7 seconds, or 'q' to quit.")

def on_press(key):
    global count
    if hasattr(key, 'char'):
        if key.char == 's':
            print("Instruction received: Screenshot will be taken in 7 seconds...")
            time.sleep(7)
            filename = f"{folder}/screenshot_{count}.png"
            screenshot = pyautogui.screenshot()
            screenshot.save(filename)

            doc.add_paragraph(f"Screenshot {count}:")
            doc.add_picture(filename, width=Inches(6))
            doc.add_paragraph("")
            print(f"✅ Screenshot {count} captured and added to Word.")
            count += 1

        elif key.char == 'q':
            doc.save("Screenshots_Report.docx")
            print("📝 Word document saved as 'Screenshots_Report.docx'. Exiting...")
            return False  # Exit listener

with keyboard.Listener(on_press=on_press) as listener:
    listener.join()
